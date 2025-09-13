import os
import faiss
import numpy as np
import pickle
import requests
import json
from sentence_transformers import SentenceTransformer
import time
import concurrent.futures
import re

# Add imports for PDF handling
try:
    import fitz  # PyMuPDF
    HAS_FITZ = True
except ImportError:
    HAS_FITZ = False

# Try to import docx for Word document processing
try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Any folder name (not full path) in this list will be skipped
EXCLUDED_DIR_NAMES = {
    'AppData', 'anaconda3', 'node_modules', '__pycache__', 'WindowsNoEditor',
    '.git', '.vscode', '.conda', '.cache', '.mamba', 'env', 'venv',
}

# Ollama API wrapper functions
def query_ollama(prompt, model="llama3", system_prompt=None, max_tokens=2000, temperature=0.7):
    """
    Send a query to the Ollama API.
    
    Args:
        prompt (str): The prompt to send to the model
        model (str): The model to use (default: llama3)
        system_prompt (str): Optional system prompt
        max_tokens (int): Maximum number of tokens to generate
        temperature (float): Sampling temperature (0.0 to 1.0)
        
    Returns:
        str: The model's response
    """
    url = "http://localhost:11434/api/generate"
    
    payload = {
        "model": model,
        "prompt": prompt,
        "stream": False,
        "max_tokens": max_tokens,
        "temperature": temperature
    }
    
    if system_prompt:
        payload["system"] = system_prompt
    
    try:
        response = requests.post(url, json=payload)
        if response.status_code == 200:
            return response.json().get("response", "")
        else:
            return f"Error: {response.status_code} - {response.text}"
    except Exception as e:
        return f"Error connecting to Ollama: {str(e)}"

def list_ollama_models():
    """List available Ollama models"""
    try:
        response = requests.get("http://localhost:11434/api/tags")
        if response.status_code == 200:
            return response.json().get("models", [])
        else:
            return []
    except:
        return []

# Document processing functions
def extract_text(file_path):
    """Extract text from a file based on its extension"""
    try:
        # Handle PDF files
        if file_path.lower().endswith('.pdf'):
            if HAS_FITZ:
                try:
                    text = ""
                    with fitz.open(file_path) as doc:
                        for page in doc:
                            text += page.get_text() + "\n"
                    return text
                except Exception as e:
                    print(f"Error extracting text from PDF {file_path}: {e}")
                    return ""
            else:
                print(f"PyMuPDF (fitz) not installed, skipping PDF file: {file_path}")
                return ""
        # Handle Word documents (.docx)
        elif file_path.lower().endswith('.docx'):
            if HAS_DOCX:
                try:
                    doc = docx.Document(file_path)
                    full_text = []
                    for para in doc.paragraphs:
                        full_text.append(para.text)
                    # Also extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                full_text.append(cell.text)
                    return '\n'.join(full_text)
                except Exception as e:
                    print(f"Error extracting text from Word document {file_path}: {e}")
                    return ""
            else:
                print(f"python-docx not installed, skipping Word file: {file_path}")
                return ""
        # Handle text files
        else:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
    except UnicodeDecodeError:
        # Try with a different encoding if UTF-8 fails
        try:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
        except Exception as e:
            print(f"Error extracting text from {file_path}: {e}")
            return ""
    except Exception as e:
        print(f"Error extracting text from {file_path}: {e}")
        return ""

def chunk_text(text, chunk_size=1000, overlap=400):
    """Split text into overlapping chunks for better semantic search"""
    if not text or not text.strip():
        return []
    
    # Detect if this is likely code by checking for common programming patterns
    likely_code = False
    code_indicators = ['def ', 'class ', 'function', 'import ', '<div', '<html', 
                      'public ', 'private ', '#include', 'module.exports',
                      'const ', 'let ', 'var ', '// ', '/* ', '"""', "'''"]
    
    for indicator in code_indicators:
        if indicator in text:
            likely_code = True
            break
    
    chunks = []
    
    # For code files, try to preserve structure by splitting on blank lines and function boundaries
    if likely_code:
        # Split by lines first
        lines = text.split('\n')
        current_chunk = []
        current_length = 0
        
        for line in lines:
            # Add the line to the current chunk
            current_chunk.append(line)
            current_length += len(line) + 1  # +1 for the newline
            
            # Check if we've reached the chunk size or hit a natural boundary
            natural_boundary = (line.strip() == '' and current_length > chunk_size / 2) or \
                              (line.startswith('def ') and current_length > chunk_size / 2) or \
                              (line.startswith('class ') and current_length > chunk_size / 2) or \
                              (line.startswith('function ') and current_length > chunk_size / 2)
            
            if current_length >= chunk_size or natural_boundary:
                # Finalize this chunk
                if current_chunk:
                    chunks.append('\n'.join(current_chunk))
                
                # Start a new chunk with overlap
                if natural_boundary:
                    # If at a natural boundary, keep some context
                    overlap_lines = min(15, len(current_chunk) // 2)  # Increased from 10 to 15
                    current_chunk = current_chunk[-overlap_lines:]
                    current_length = sum(len(line) + 1 for line in current_chunk)
                else:
                    # Otherwise, keep more of the previous lines for context/overlap
                    overlap_lines = min(20, len(current_chunk) // 2)  # Increased from 10 to 20
                    current_chunk = current_chunk[-overlap_lines:]
                    current_length = sum(len(line) + 1 for line in current_chunk)
        
        # Add the last chunk if it's not empty
        if current_chunk:
            chunks.append('\n'.join(current_chunk))
    else:
        # For normal text, use a semantic-aware approach that tries to break at sentence boundaries
        sentences = re.split(r'(?<=[.!?])\s+', text)
        current_chunk = []
        current_length = 0
        
        for sentence in sentences:
            current_chunk.append(sentence)
            current_length += len(sentence) + 1  # +1 for space
            
            # Split when we exceed chunk size, preferring sentence boundaries
            if current_length >= chunk_size:
                chunks.append(' '.join(current_chunk))
                
                # Keep some sentences for overlap
                overlap_sentences = max(3, len(current_chunk) // 3)  # At least 3 sentences overlap
                current_chunk = current_chunk[-overlap_sentences:]
                current_length = sum(len(s) + 1 for s in current_chunk)
        
        # Add the last chunk if not empty
        if current_chunk:
            chunks.append(' '.join(current_chunk))
    
    return chunks

# Vector embedding functions
_model = None

def get_embedding_model(model_name="sentence-transformers/all-mpnet-base-v2"):
    """Get a singleton embedding model instance"""
    global _model
    if _model is None:
        _model = SentenceTransformer(model_name)
    return _model

def embed_text(text, model=None):
    """Generate embeddings for a piece of text"""
    if not model:
        model = get_embedding_model()
    
    if isinstance(text, str):
        return model.encode([text], normalize_embeddings=True)[0]
    else:
        return model.encode(text, normalize_embeddings=True)

# Document indexing
def process_document(path, model=None, use_cache=True, cache_dir="index/cache"):
    """
    Process a single document and return a list of (embedding, metadata) tuples,
    one for each chunk.
    """
    if use_cache and not os.path.exists(cache_dir):
        os.makedirs(cache_dir, exist_ok=True)

    # Generate cache key
    cache_key = None
    if use_cache:
        try:
            file_stat = os.stat(path)
            modified_time = file_stat.st_mtime
            file_size = file_stat.st_size
            cache_key = f"{os.path.abspath(path)}_{modified_time}_{file_size}"
            cache_key = cache_key.replace('\\', '_').replace('/', '_').replace(':', '_')
            cache_file = os.path.join(cache_dir, f"{cache_key}.pkl")

            # Load from cache
            if os.path.exists(cache_file):
                try:
                    with open(cache_file, "rb") as f:
                        cached_data = pickle.load(f)
                        print(f"Cache hit for {path}")
                        return cached_data
                except Exception as e:
                    print(f"Error loading cache for {path}: {e}")
        except Exception as e:
            print(f"Cache key generation error for {path}: {e}")

    try:
        if model is None:
            model = get_embedding_model()

        text = extract_text(path)
        if not text or not text.strip():
            print(f"No content extracted from {path}")
            return []

        chunks = chunk_text(text)
        if not chunks:
            print(f"No chunks generated for {path}")
            return []

        chunk_data = []
        for chunk in chunks:
            embedding = model.encode([chunk], normalize_embeddings=True)[0]
            metadata = {
                "chunk": chunk,
                "filename": os.path.basename(path),
                "path": path
            }
            chunk_data.append((embedding, metadata))

        # Save to cache
        if use_cache and cache_key:
            try:
                with open(cache_file, "wb") as f:
                    pickle.dump(chunk_data, f)
                print(f"Saved to cache: {cache_file}")
            except Exception as e:
                print(f"Error saving to cache for {path}: {e}")

        return chunk_data

    except Exception as e:
        print(f"Error processing {path}: {e}")
        return []


def build_document_index(file_paths, index_path="index/faiss_index.bin", metadata_path="index/metadata.pkl",
                         progress_callback=None, use_parallel=True, max_workers=None):
    """
    Build a FAISS vector index from document chunks (not whole documents).
    Each chunk is indexed separately with its own metadata.

    Args:
        file_paths: List of file paths to index
        index_path: Path to save the FAISS index
        metadata_path: Path to save the metadata
        progress_callback: Optional callback function(progress_float, message)
        use_parallel: Whether to use parallel processing
        max_workers: Number of worker threads (None = auto)

    Returns:
        bool: True if successful
    """
    if not file_paths:
        if progress_callback:
            progress_callback(1.0, "No files to index")
        return False

    print(f"Building index for {len(file_paths)} files")
    model = get_embedding_model()

    os.makedirs(os.path.dirname(index_path), exist_ok=True)

    # Load existing index and metadata (if any)
    existing_index = None
    existing_metadata = []
    already_indexed_paths = set()

    if os.path.exists(index_path) and os.path.exists(metadata_path):
        try:
            existing_index = faiss.read_index(index_path)
            with open(metadata_path, "rb") as f:
                existing_metadata = pickle.load(f)
            already_indexed_paths = set(m["path"] for m in existing_metadata)
            print(f"Found existing index with {len(existing_metadata)} chunks")
        except Exception as e:
            print(f"Error loading existing index or metadata: {e}")

    # Filter new files
    new_files = [path for path in file_paths if path not in already_indexed_paths]
    if not new_files:
        if progress_callback:
            progress_callback(1.0, "No new files to index")
        return True

    print(f"Processing {len(new_files)} new files")
    all_embeddings = []
    all_metadata = []

    # Progress function
    def update_progress(completed, total):
        if progress_callback:
            progress_callback((completed / total) * 0.9, f"Processed {completed}/{total} files")

    # Process documents (parallel or sequential)
    if use_parallel and len(new_files) > 1:
        if not max_workers:
            max_workers = min(os.cpu_count() or 2, 8)
        print(f"Using {max_workers} parallel workers")

        with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = {executor.submit(process_document, path, model): path for path in new_files}
            for i, future in enumerate(concurrent.futures.as_completed(futures)):
                path = futures[future]
                try:
                    chunk_data = future.result()
                    for emb, meta in chunk_data:
                        all_embeddings.append(emb)
                        all_metadata.append(meta)
                except Exception as e:
                    print(f"Error processing {path}: {e}")
                update_progress(i + 1, len(new_files))
    else:
        print("Processing files sequentially")
        for i, path in enumerate(new_files):
            print(f"Processing {i + 1}/{len(new_files)}: {path}")
            try:
                chunk_data = process_document(path, model)
                for emb, meta in chunk_data:
                    all_embeddings.append(emb)
                    all_metadata.append(meta)
            except Exception as e:
                print(f"Error processing {path}: {e}")
            update_progress(i + 1, len(new_files))

    if not all_embeddings:
        print("No valid embeddings were generated.")
        if progress_callback:
            progress_callback(1.0, "No valid content found.")
        return False

    # Final progress update
    if progress_callback:
        progress_callback(0.95, "Building FAISS index")

    try:
        embedding_matrix = np.vstack(all_embeddings)
        dim = embedding_matrix.shape[1]
        print(f"Creating FAISS index with {embedding_matrix.shape[0]} chunks (dim={dim})")

        if existing_index and existing_index.d == dim:
            print(f"Appending to existing index (currently has {existing_index.ntotal} vectors)")
            existing_index.add(embedding_matrix)
            combined_metadata = existing_metadata + all_metadata
            faiss.write_index(existing_index, index_path)
            with open(metadata_path, "wb") as f:
                pickle.dump(combined_metadata, f)
        else:
            print("Creating new FAISS index")
            index = faiss.IndexFlatIP(dim)  # cosine-compatible
            index.add(embedding_matrix)
            faiss.write_index(index, index_path)
            with open(metadata_path, "wb") as f:
                pickle.dump(all_metadata, f)

        if progress_callback:
            progress_callback(1.0, "Index built successfully")
        print("✅ Indexing complete")
        return True

    except Exception as e:
        print(f"Error building index: {e}")
        if progress_callback:
            progress_callback(1.0, f"Error: {e}")
        return False

# Document search
def search_documents(query, top_k=5, index_path="index/faiss_index.bin", metadata_path="index/metadata.pkl", 
                      threshold=0.1):  # Lowered threshold from 0.3 to 0.1
    """
    Search indexed document chunks using semantic similarity.

    Args:
        query (str): The user's search query
        top_k (int): Number of top results to return
        index_path (str): Path to the FAISS index file
        metadata_path (str): Path to the metadata file
        threshold (float): Similarity threshold (lower = more results but might be less relevant)

    Returns:
        list: A list of result dicts, or a dict with "error"
    """
    try:
        if not os.path.exists(index_path) or not os.path.exists(metadata_path):
            return {"error": "Index not found. Please build the index first."}

        # Load index and metadata
        index = faiss.read_index(index_path)
        with open(metadata_path, "rb") as f:
            metadata = pickle.load(f)

        # Generate query embedding
        model = get_embedding_model()
        query_vec = model.encode([query], normalize_embeddings=True)

        # Perform the FAISS search - retrieve more results initially
        expanded_k = min(top_k * 3, len(metadata))  # Get more results to filter
        distances, indices = index.search(query_vec, expanded_k)

        results = []
        for i, idx in enumerate(indices[0]):
            if idx < 0 or idx >= len(metadata):
                continue

            similarity_score = float(distances[0][i])
            if similarity_score < threshold:  # Using the lower threshold parameter
                continue

            entry = metadata[idx]
            results.append({
                "score": similarity_score,
                "filename": entry.get("filename", ""),
                "file_path": entry.get("path", ""),
                "snippet": entry.get("chunk", "")
            })

        # Return top_k results after filtering
        return results[:top_k]

    except Exception as e:
        return {"error": f"Search error: {str(e)}"}


def scan_directory(directory, extensions=None):
    """
    Scan a directory for files to index
    
    Args:
        directory: Directory path to scan
        extensions: List of file extensions to include (None = all)
        
    Returns:
        list: List of file paths
    """
    if extensions is None:
        extensions = ['.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.csv', '.pdf']
    
    file_paths = []
    
    # Convert extensions to lowercase for case-insensitive matching
    extensions = [ext.lower() for ext in extensions]
    
    try:
        for root, dirs, files in os.walk(directory):
            # Skip excluded directories
            dirs[:] = [d for d in dirs if d not in EXCLUDED_DIR_NAMES]
            
            for file in files:
                try:
                    if not extensions or any(file.lower().endswith(ext) for ext in extensions):
                        file_path = os.path.join(root, file)
                        # Verify file is accessible and not too large
                        try:
                            if os.path.isfile(file_path) and os.access(file_path, os.R_OK):
                                # Skip files larger than 50MB to prevent memory issues
                                if os.path.getsize(file_path) < 50 * 1024 * 1024:
                                    file_paths.append(file_path)
                                else:
                                    print(f"Skipping large file: {file_path}")
                        except Exception as e:
                            print(f"Error accessing file {file_path}: {e}")
                except Exception as e:
                    print(f"Error processing file entry: {e}")
    except Exception as e:
        print(f"Error walking directory {directory}: {e}")
    
    return file_paths

# Document Q&A with Ollama
def answer_question(question, context, model="llama3"):
    """
    Answer a question using Ollama with context from document search
    
    Args:
        question: The user's question
        context: Document context (from search)
        model: Ollama model to use
        
    Returns:
        str: The answer
    """
    # Format search results into context
    context_text = ""
    for i, result in enumerate(context):
        context_text += f"\nDocument {i+1}: {result['filename']}\n"
        context_text += f"Content: {result['snippet']}\n"
    
    # Create a prompt with context
    prompt = f"""Based on the following documents, please answer this question:

Question: {question}

Context:
{context_text}

Answer:"""
    
    # System prompt to guide the model
    system_prompt = """You are a helpful assistant that answers questions based on the provided document context.
If the answer is not in the documents, say 'I don't have enough information to answer this question.'
Always cite the specific document you used to answer the question."""
    
    # Query Ollama
    response = query_ollama(prompt, model=model, system_prompt=system_prompt)
    return response 