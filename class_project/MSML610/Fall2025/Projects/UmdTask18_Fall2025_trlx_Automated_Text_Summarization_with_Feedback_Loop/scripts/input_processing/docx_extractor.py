"""
DOCX Text Extractor

Extracts and cleans text from DOCX files using python-docx.
Preserves paragraph structure while removing headers/footers.
"""

from pathlib import Path
from typing import Optional
import logging

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not installed. Install with: pip install python-docx")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DOCXExtractor:
    """Extract text from DOCX files."""
    
    def __init__(self):
        """Initialize DOCX extractor."""
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX extraction. "
                "Install with: pip install python-docx"
            )
    
    def extract_from_docx(self, docx_path: str) -> Optional[str]:
        """
        Extract text from DOCX file.
        
        Args:
            docx_path: Path to DOCX file
        
        Returns:
            Extracted text or None if extraction fails
        """
        docx_path = Path(docx_path)
        
        if not docx_path.exists():
            logger.error(f"DOCX file not found: {docx_path}")
            return None
        
        try:
            logger.info(f"Extracting text from DOCX: {docx_path}")
            
            doc = Document(docx_path)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:  # Skip empty paragraphs
                    paragraphs.append(text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)
            
            if paragraphs:
                full_text = "\n\n".join(paragraphs)
                logger.info(f"Successfully extracted {len(full_text)} characters from {len(paragraphs)} paragraphs")
                return full_text
            else:
                logger.warning(f"No text extracted from {docx_path}")
                return None
                
        except Exception as e:
            logger.error(f"Error extracting from DOCX {docx_path}: {e}")
            return None
    
    def extract_with_metadata(self, docx_path: str) -> dict:
        """
        Extract text with metadata.
        
        Args:
            docx_path: Path to DOCX file
        
        Returns:
            Dictionary with text and metadata
        """
        docx_path = Path(docx_path)
        
        try:
            doc = Document(docx_path)
            
            # Extract text
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)
            
            full_text = "\n\n".join(paragraphs) if paragraphs else None
            
            # Extract core properties
            core_props = doc.core_properties
            
            return {
                "file_path": str(docx_path),
                "text": full_text,
                "num_paragraphs": len(paragraphs),
                "title": core_props.title if hasattr(core_props, 'title') else None,
                "author": core_props.author if hasattr(core_props, 'author') else None,
                "subject": core_props.subject if hasattr(core_props, 'subject') else None,
                "created": str(core_props.created) if hasattr(core_props, 'created') and core_props.created else None,
                "modified": str(core_props.modified) if hasattr(core_props, 'modified') and core_props.modified else None,
                "success": full_text is not None
            }
            
        except Exception as e:
            logger.error(f"Error extracting from DOCX {docx_path}: {e}")
            return {
                "file_path": str(docx_path),
                "text": None,
                "num_paragraphs": 0,
                "title": None,
                "author": None,
                "subject": None,
                "created": None,
                "modified": None,
                "success": False,
                "error": str(e)
            }


def extract_from_docx(docx_path: str) -> Optional[str]:
    """
    Convenience function to extract from DOCX.
    
    Args:
        docx_path: Path to DOCX file
    
    Returns:
        Extracted text or None
    """
    extractor = DOCXExtractor()
    return extractor.extract_from_docx(docx_path)


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1:
        docx_file = sys.argv[1]
        
        extractor = DOCXExtractor()
        result = extractor.extract_with_metadata(docx_file)
        
        if result["success"]:
            print(f"File: {result['file_path']}")
            print(f"Paragraphs: {result['num_paragraphs']}")
            print(f"Title: {result['title']}")
            print(f"Author: {result['author']}")
            print(f"Created: {result['created']}")
            print(f"\nText preview (first 500 chars):")
            print(result['text'][:500])
            print(f"\n... ({len(result['text'])} total characters)")
        else:
            print(f"Extraction failed: {result.get('error', 'Unknown error')}")
    else:
        print("Usage: python docx_extractor.py <docx_file>")
